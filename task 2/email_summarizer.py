import os
import json
import streamlit as st
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from fpdf import FPDF

# ----------------------------
# Setup
# ----------------------------
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

if not GOOGLE_API_KEY:
    st.error("❌ Missing GOOGLE_API_KEY in .env file")
    st.stop()

genai.configure(api_key=GOOGLE_API_KEY)

st.set_page_config(page_title="📨 Email Summarizer Agent", layout="wide")

# ----------------------------
# Helper Functions
# ----------------------------
def load_emails_from_file(uploaded_file):
    """Load emails from various file formats."""
    ext = os.path.splitext(uploaded_file.name)[-1].lower()

    try:
        if ext == ".json":
            return json.load(uploaded_file)
        elif ext == ".txt":
            text = uploaded_file.read().decode("utf-8")
            return [{"sender": "unknown", "subject": "Text Upload", "body": text}]
        elif ext == ".pdf":
            import fitz  # PyMuPDF
            pdf = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            text = "\n".join(page.get_text() for page in pdf)
            return [{"sender": "unknown", "subject": "PDF Upload", "body": text}]
        elif ext == ".docx":
            doc = Document(uploaded_file)
            text = "\n".join(p.text for p in doc.paragraphs)
            return [{"sender": "unknown", "subject": "Word Upload", "body": text}]
        else:
            st.warning("⚠️ Unsupported file type.")
            return []
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return []

def summarize_emails(emails, model_choice):
    """Use Gemini to summarize the given emails."""
    if not emails:
        return "No emails found."

    model = genai.GenerativeModel(model_choice)

    prompt = (
        "You are an AI assistant that summarizes emails.\n"
        "Summarize each email with key points and tone.\n"
        "Output format:\n\n"
        "1. Subject: ...\n   From: ...\n   Summary: ...\n\n"
        f"Emails:\n{json.dumps(emails, indent=2)}"
    )

    response = model.generate_content(prompt)
    return response.text

def save_summary_as_docx(summary_text, filename="email_summary.docx"):
    """Save the summary as a DOCX file."""
    doc = Document()
    doc.add_heading("📨 AI Email Summary", level=1)
    doc.add_paragraph(summary_text)
    doc.save(filename)
    return filename

def save_summary_as_pdf(summary_text, filename="email_summary.pdf"):
    """Save the summary as a PDF file."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in summary_text.split("\n"):
        pdf.multi_cell(0, 10, line)
    pdf.output(filename)
    return filename

# ----------------------------
# UI Layout
# ----------------------------
st.markdown(
    "<h1 style='text-align:center; color:#00b7ff;'>📨 Email Summarizer Agent</h1>",
    unsafe_allow_html=True,
)
st.markdown("---")

# Mode Selection
mode = st.radio("🧭 Choose Mode:", ["Dashboard Mode", "Classic Mode"], horizontal=True)

# Model Selection
model_choice = st.selectbox(
    "🧠 Choose Gemini Model",
    ["models/gemini-2.5-flash", "models/gemini-2.5-pro", "models/gemini-flash-latest"],
)

# Toggle for sample vs upload
use_sample_data = st.toggle("🧩 Use Sample Data For Test Purpose")

# File uploader (only shown if not using sample)
uploaded_file = None
if not use_sample_data:
    uploaded_file = st.file_uploader(
        "📂 Upload your emails file (JSON, TXT, PDF, DOCX)",
        type=["json", "txt", "pdf", "docx"],
    )

# Load data
emails = []
if use_sample_data:
    if os.path.exists("emails.json"):
        with open("emails.json", "r", encoding="utf-8") as f:
            emails = json.load(f)
        st.success("✅ Sample data loaded successfully.")
    else:
        st.error("⚠️ No sample file found (emails.json). Please upload your own file.")
else:
    if uploaded_file:
        emails = load_emails_from_file(uploaded_file)
    else:
        st.info("📥 Upload a file or switch on 'Use Sample Data' to test.")

# ----------------------------
# Dashboard Mode
# ----------------------------
if mode == "Dashboard Mode":
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("📬 Loaded Emails")
        if emails:
            for i, email in enumerate(emails, 1):
                st.markdown(
                    f"**{i}. From:** {email.get('sender', 'N/A')}  \n"
                    f"**Subject:** {email.get('subject', 'N/A')}  \n"
                    f"📄 *{email.get('body', '')[:300]}...*"
                )
                st.markdown("---")
        else:
            st.warning("No emails loaded.")

    with col2:
        st.subheader("🧠 AI Summary")

        btn_label = "🚀 Summarize Sample Data" if use_sample_data else "🚀 Summarize Uploaded Data"
        if st.button(btn_label, use_container_width=True):
            with st.spinner("Summarizing emails with Gemini..."):
                summary = summarize_emails(emails, model_choice)

            st.success("✅ Summarization Complete!")
            st.text_area("Summary Output", summary, height=400)

            # Downloads
            st.download_button(
                "💾 Download as JSON",
                json.dumps({"summary": summary}, indent=2),
                "email_summary.json",
                "application/json",
            )

            docx_path = save_summary_as_docx(summary)
            pdf_path = save_summary_as_pdf(summary)

            with open(docx_path, "rb") as f:
                st.download_button(
                    "📘 Download as Word (.docx)", f, file_name="email_summary.docx"
                )
            with open(pdf_path, "rb") as f:
                st.download_button(
                    "📄 Download as PDF", f, file_name="email_summary.pdf"
                )

# ----------------------------
# Classic Mode
# ----------------------------
else:
    if emails:
        st.subheader("📬 Loaded Emails")
        for email in emails:
            st.markdown(
                f"**From:** {email.get('sender', 'N/A')}  \n"
                f"**Subject:** {email.get('subject', 'N/A')}  \n"
                f"📄 *{email.get('body', '')}*"
            )
            st.markdown("---")

        btn_label = "🚀 Summarize Sample Data" if use_sample_data else "🚀 Summarize Uploaded Data"
        if st.button(btn_label):
            with st.spinner("Summarizing emails..."):
                summary = summarize_emails(emails, model_choice)
            st.success("✅ Summarization Complete!")
            st.markdown("### 🧠 AI Summary")
            st.markdown(f"```markdown\n{summary}\n```")

            # Downloads
            st.download_button(
                "💾 Download as JSON",
                json.dumps({"summary": summary}, indent=2),
                "email_summary.json",
                "application/json",
            )

            docx_path = save_summary_as_docx(summary)
            pdf_path = save_summary_as_pdf(summary)

            with open(docx_path, "rb") as f:
                st.download_button(
                    "📘 Download as Word (.docx)", f, file_name="email_summary.docx"
                )
            with open(pdf_path, "rb") as f:
                st.download_button(
                    "📄 Download as PDF", f, file_name="email_summary.pdf"
                )
    else:
        st.warning("No email data found. Please upload or enable sample data.")
